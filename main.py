import streamlit as st
import yaml
import bcrypt
import os
from pathlib import Path
import asyncio
from typing import List, Dict, Any, Optional
import io
import re # Added for link extraction
from urllib.parse import urlparse, urljoin # Added for link processing
import pandas as pd # Import pandas

try:
    import fitz  # PyMuPDF
except ImportError:
    # This will be caught at runtime if not installed. 
    # User will need to add PyMuPDF to requirements.txt and rebuild.
    pass 
try:
    from docx import Document # python-docx
except ImportError:
    # This will be caught at runtime if not installed.
    # User will need to add python-docx to requirements.txt and rebuild.
    pass

from src.config import (
    DEBUG,
    LOG_LEVEL,
    OUTPUT_FORMAT,
    OUTPUT_DIR,
    DEFAULT_PROMPTS,
    USERS_CONFIG_PATH,
    SYSTEM_PROMPT,
    OPENROUTER_PRIMARY_MODEL # Import the primary model default
)
from src.openrouter import OpenRouterClient
from src.firecrawl_client import FirecrawlClient
from src.audit_logger import log_audit_event # Added for Task 5

# Initialize clients
@st.cache_resource
def init_clients():
    openrouter_client = OpenRouterClient()
    firecrawl_client = FirecrawlClient(
        redis_url=os.getenv("REDIS_URL")
    )
    return openrouter_client, firecrawl_client

def load_users() -> Dict[str, Any]:
    """Load user data from YAML file."""
    if not os.path.exists(USERS_CONFIG_PATH):
        # If no users.yaml, create it with default admin/researcher from init_users.py logic
        # This is a fallback, ideally init_users.py should be run once.
        from src.init_users import init_users as initialize_system_users
        try:
            initialize_system_users()
            st.info("User configuration file not found. Initialized with default users.")
        except Exception as e:
            st.error(f"Failed to initialize default users: {e}")
            return {}
            
    with open(USERS_CONFIG_PATH, 'r') as f:
        users = yaml.safe_load(f)
        if users is None: # Handle empty users.yaml file
            st.warning("User configuration file is empty. Please run user initialization or sign up.")
            return {}
        return users

def save_users(users_data: Dict[str, Any]) -> bool:
    try:
        with open(USERS_CONFIG_PATH, 'w') as f:
            yaml.dump(users_data, f, default_flow_style=False, sort_keys=False)
        return True
    except Exception as e:
        st.error(f"Failed to save user data: {e}")
        return False

def verify_password(plain_password: str, hashed_password: str) -> bool:
    """Verify a password against its hash."""
    return bcrypt.checkpw(
        plain_password.encode(),
        hashed_password.encode()
    )

def hash_password(password: str) -> str:
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

# Helper functions for document processing
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extracts text from PDF file bytes."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        st.error(f"Error processing PDF: {e}")
        return ""

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extracts text from DOCX file bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error processing DOCX: {e}")
        return ""

def extract_text_from_txt_md(file_bytes: bytes, encoding='utf-8') -> str:
    """Extracts text from TXT or MD file bytes."""
    try:
        return file_bytes.decode(encoding)
    except UnicodeDecodeError:
        try:
            return file_bytes.decode('latin-1') # Fallback encoding
        except Exception as e:
            st.error(f"Error decoding text file with fallback: {e}")
            return ""
    except Exception as e:
        st.error(f"Error processing text/markdown file: {e}")
        return ""
# End of helper functions

async def process_urls(urls: List[str], client: FirecrawlClient) -> List[Dict[str, Any]]:
    """Process a list of specific URLs and return a list of dicts with url, content/error."""
    results = await client.scrape_multiple_urls(urls)
    
    processed_results = []
    for result in results:
        url = result.get("metadata", {}).get("url", result.get("url", "unknown URL"))
        if result.get("success", False):
            content_data = result.get("data", {}).get("content", "")
            if not content_data:
                 content_data = result.get("content", "")
            processed_results.append({"url": url, "content": content_data, "status": "success"})
        else:
            error_message = result.get("error", "Unknown error")
            processed_results.append({"url": url, "error": error_message, "status": "failed"})
            
    return processed_results

async def perform_web_research(query: str, client: OpenRouterClient) -> str:
    """Perform web research on a given query."""
    # Generate search queries
    queries = await client.generate_serp_queries(query)
    
    # Simulate search results (replace with actual search implementation)
    search_results = {
        "data": [
            {"content": "Sample content 1", "url": "https://example.com/1"},
            {"content": "Sample content 2", "url": "https://example.com/2"}
        ]
    }
    
    # Process search results
    results = await client.process_serp_result(query, search_results)
    
    # Write final report
    report = await client.write_final_report(
        query,
        results["learnings"],
        [item["url"] for item in search_results["data"]]
    )
    
    return report

# --- Add Model Definitions Here ---
# Define model choices based on user query and web search results
MODEL_OPTIONS = {
    "Mistral Medium 3": "mistralai/mistral-medium-3",
    "Google Gemini 2.5 Pro": "google/gemini-2.5-pro-preview",
    "OpenAI o3": "openai/o3",
    "OpenAI GPT-4.1": "openai/gpt-4.1",
    "Claude 3.7 Sonnet (thinking)": "anthropic/claude-3.7-sonnet:thinking",
    "DeepSeek R1T Chimera": "tngtech/deepseek-r1t-chimera:free",
}
MODEL_DISPLAY_NAMES = list(MODEL_OPTIONS.keys())
MODEL_IDENTIFIERS = list(MODEL_OPTIONS.values())
# --- End Model Definitions ---

# --- NEW: Crawl and Scrape Function --- 
async def crawl_and_scrape_site(start_url: str, limit: int, client: FirecrawlClient) -> List[Dict[str, Any]]:
    """Crawls a website starting from start_url, scraping content and following same-domain links up to a limit."""
    if not start_url or not client.validate_url(start_url):
        st.warning(f"Invalid start URL provided for crawl: {start_url}")
        return []

    if limit <= 0:
        st.warning("Crawl limit must be greater than 0.")
        return []

    base_domain = urlparse(start_url).netloc
    urls_to_scrape = {start_url}
    visited_urls = set()
    scraped_data_list = []
    scrape_count = 0

    # Modified regex to use r"..." string format
    link_regex = re.compile(r'href\s*=\s*[\'"]([^\'"]+)[\'"]', re.IGNORECASE)

    st.info(f"Starting crawl from {start_url} (Domain: {base_domain}), limit: {limit} pages.")

    while urls_to_scrape and scrape_count < limit:
        current_url = urls_to_scrape.pop()

        if current_url in visited_urls:
            continue

        visited_urls.add(current_url)
        scrape_count += 1
        st.write(f"Crawling [{scrape_count}/{limit}]: {current_url}") # Show progress

        try:
            # Use the single scrape_url method from the client
            scraped_result_dict = await client.scrape_url(current_url)
            
            # Extract markdown content for storage/AI processing
            markdown_content = scraped_result_dict.get("data", {}).get("content", "") 
            # Extract HTML content specifically for link extraction
            html_for_links = scraped_result_dict.get("data", {}).get("html_content", "")
            
            # Check for errors reported by the client
            client_error = scraped_result_dict.get("error")
            if client_error:
                st.error(f"Scraper client returned error for {current_url}: {client_error}")
                scraped_data_list.append({"url": current_url, "error": client_error, "status": "failed"})
                log_audit_event(
                    username=st.session_state.get('username', 'SYSTEM'),
                    role=st.session_state.get('role', 'N/A'),
                    action="CRAWL_SCRAPE_CLIENT_ERROR", # Distinguish from API error
                    details=f"Firecrawl client processing error for {current_url}: {client_error}",
                    links=[current_url]
                )
                continue # Move to the next URL

            if markdown_content:
                # Store the markdown content
                scraped_data_list.append({"url": current_url, "content": markdown_content, "status": "success"})
                
                # Find and process links using HTML content if available, otherwise fallback to markdown
                link_source_content = html_for_links if html_for_links else markdown_content
                if scrape_count < limit and link_source_content:
                    found_links_count = 0
                    potential_links = link_regex.findall(link_source_content)
                    for link in potential_links:
                        try:
                            absolute_link = urljoin(current_url, link.strip())
                            parsed_link = urlparse(absolute_link)
                            
                            # Basic validation and domain check
                            if parsed_link.scheme in ['http', 'https'] and parsed_link.netloc == base_domain:
                                if absolute_link not in visited_urls and absolute_link not in urls_to_scrape:
                                    urls_to_scrape.add(absolute_link)
                                    found_links_count += 1
                        except Exception as link_e:
                            # Ignore errors parsing/resolving individual links
                            # print(f" Minor error processing link '{link}': {link_e}")
                            pass # Be less verbose in UI
                    # st.write(f"  Found {found_links_count} new links on {current_url}") # Optional debug
            else: # No markdown content extracted
                error_msg = f"No primary (markdown) content extracted from: {current_url}"
                scraped_data_list.append({"url": current_url, "error": error_msg, "status": "no_content"})
                st.warning(error_msg)
        
        except Exception as e: # Catch exceptions from scrape_url call itself or processing
            error_msg = f"Error scraping {current_url}: {str(e)}"
            st.error(error_msg)
            scraped_data_list.append({"url": current_url, "error": error_msg, "status": "failed"})
            # Log this significant error
            log_audit_event(
                username=st.session_state.get('username', 'SYSTEM'), # Get username if available
                role=st.session_state.get('role', 'N/A'),
                action="CRAWL_SCRAPE_FAILURE",
                details=error_msg,
                links=[current_url]
            )

    if urls_to_scrape:
        st.warning(f"Crawl limit ({limit}) reached, {len(urls_to_scrape)} URLs remaining in queue.")
    
    successful_scrape_count = sum(1 for item in scraped_data_list if item.get('status') == 'success' and item.get('content'))
    st.info(f"Crawl finished. Attempted {len(scraped_data_list)} pages, successfully scraped {successful_scrape_count} with content.")
    return scraped_data_list
# --- End Crawl Function --- 

# --- Add Function to Parse Log Line ---
def parse_log_line(line: str) -> Optional[Dict[str, str]]:
    """Parses a single line from the audit log file."""
    parts = line.strip().split(' | ')
    if len(parts) < 4: # Basic check for minimum parts (Timestamp, User, Role, Action)
        return None
    
    log_entry = {
        "Timestamp": parts[0],
        "Username": "N/A",
        "Role": "N/A",
        "Action": "N/A",
        "Model": "N/A",
        "Links": "N/A",
        "Details": "N/A"
    }
    
    for part in parts[1:]: # Skip timestamp
        try:
            key, value = part.split(': ', 1)
            # Map keys, handling potential variations
            if key == "USER": log_entry["Username"] = value
            elif key == "ROLE": log_entry["Role"] = value
            elif key == "ACTION": log_entry["Action"] = value
            elif key == "MODEL": log_entry["Model"] = value
            elif key == "LINKS": log_entry["Links"] = value
            elif key == "DETAILS": log_entry["Details"] = value
        except ValueError:
             # If a part doesn't split correctly, add it to Details as unparsed info
             log_entry["Details"] = f"{log_entry.get('Details', '')} [Unparsed: {part}]"
            
    return log_entry
# --- End Function ---

async def main():
    st.set_page_config(
        page_title="AI Research Agent",
        page_icon="🤖",
        layout="wide"
    )

    # Initialize session state for system_prompt if not already set
    if "system_prompt" not in st.session_state:
        st.session_state.system_prompt = SYSTEM_PROMPT # Global default
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    if "username" not in st.session_state:
        st.session_state.username = None
    if "show_signup" not in st.session_state:
        st.session_state.show_signup = False
    if "processed_documents_content" not in st.session_state: 
        st.session_state.processed_documents_content = []
    if "last_uploaded_file_details" not in st.session_state: # For comparing if files changed
        st.session_state.last_uploaded_file_details = []
    if "unified_report_content" not in st.session_state: 
        st.session_state.unified_report_content = ""
    if "scraped_web_content" not in st.session_state: # Added for this task (Task 3)
        st.session_state.scraped_web_content = []
    if "crawled_web_content" not in st.session_state:
        st.session_state.crawled_web_content = []

    # Sidebar
    with st.sidebar:
        st.title("Access Panel")
        
        if not st.session_state.authenticated:
            if st.session_state.show_signup:
                st.subheader("Create Account")
                new_username = st.text_input("New Username", key="signup_username")
                new_password = st.text_input("New Password", type="password", key="signup_password")
                confirm_password = st.text_input("Confirm Password", type="password", key="signup_confirm_password")
                
                if st.button("Sign Up", key="signup_submit_button"):
                    if not new_username or not new_password:
                        st.error("Username and password cannot be empty.")
                        log_audit_event(new_username or "ANONYMOUS", "N/A", "USER_SIGNUP_FAILURE", "Attempted signup with empty username/password.")
                    elif new_password != confirm_password:
                        st.error("Passwords do not match.")
                        log_audit_event(new_username, "N/A", "USER_SIGNUP_FAILURE", "Attempted signup with non-matching passwords.")
                    else:
                        users = load_users()
                        if new_username in users:
                            st.error("Username already exists.")
                            log_audit_event(new_username, "N/A", "USER_SIGNUP_FAILURE", f"Attempted signup with existing username: {new_username}.")
                        else:
                            users[new_username] = {
                                "password": hash_password(new_password),
                                "role": "researcher", # Default role
                                "system_prompt": DEFAULT_PROMPTS.get("researcher", SYSTEM_PROMPT)
                            }
                            if save_users(users):
                                st.success("Account created successfully! Please log in.")
                                log_audit_event(new_username, "researcher", "USER_SIGNUP_SUCCESS", f"New user account created: {new_username}")
                                st.session_state.show_signup = False # Switch back to login
                            else:
                                st.error("Failed to create account. Please try again.")
                                log_audit_event(new_username, "researcher", "USER_SIGNUP_FAILURE", f"Failed to save new user account: {new_username} after validation.")
                
                if st.button("Back to Login", key="back_to_login_button"):
                    st.session_state.show_signup = False
                    st.rerun()
            else:
                st.subheader("Login")
                username_input = st.text_input("Username", key="login_username")
                password_input = st.text_input("Password", type="password", key="login_password")
                
                if st.button("Login", key="login_button"):
                    users = load_users()
                    user_data = users.get(username_input, {})
                    if user_data and verify_password(password_input, user_data.get("password", "")):
                        st.session_state.authenticated = True
                        st.session_state.username = username_input
                        # Load user-specific prompt, fallback to role-based, then global default
                        user_specific_prompt = user_data.get("system_prompt")
                        current_role = user_data.get("role", "researcher") # Default role if not specified
                        st.session_state.role = current_role # Store role in session state

                        if not user_specific_prompt:
                            user_specific_prompt = DEFAULT_PROMPTS.get(current_role, SYSTEM_PROMPT)
                        
                        st.session_state.system_prompt = user_specific_prompt
                        st.success("Login successful!")
                        log_audit_event(username_input, current_role, "USER_LOGIN_SUCCESS", f"User {username_input} logged in successfully.")
                        st.rerun()
                    else:
                        st.error("Invalid username or password")
                        log_audit_event(username_input or "UNKNOWN_USER", "N/A", "USER_LOGIN_FAILURE", f"Failed login attempt for username: '{username_input}'.")
                
                if st.button("Create Account", key="show_signup_button"):
                    st.session_state.show_signup = True
                    st.rerun()
        else:
            st.write(f"Logged in as: {st.session_state.username}")
            if st.button("Logout", key="logout_button"):
                logged_out_username = st.session_state.username
                logged_out_role = st.session_state.get("role", "N/A")
                st.session_state.authenticated = False
                st.session_state.username = None
                st.session_state.role = None # Clear role on logout
                st.session_state.system_prompt = SYSTEM_PROMPT # Reset to global default on logout
                st.session_state.show_signup = False # Reset signup view on logout
                log_audit_event(logged_out_username, logged_out_role, "USER_LOGOUT", f"User {logged_out_username} logged out.")
                st.rerun()

            st.markdown("---") # Separator
            st.subheader("Session System Prompt")
            # Allow editing of the session system prompt
            new_prompt = st.text_area(
                "Edit current session's system prompt:", 
                value=st.session_state.system_prompt, 
                height=300,
                key="session_prompt_editor"
            )
            if new_prompt != st.session_state.system_prompt:
                st.session_state.system_prompt = new_prompt
                st.success("Session system prompt updated.")
                log_audit_event(
                    username=st.session_state.username, 
                    role=st.session_state.get("role", "N/A"), 
                    action="SESSION_PROMPT_UPDATED", 
                    details=f"User updated session prompt. New prompt: '{new_prompt}'"
                )

    # Main content area
    if st.session_state.authenticated:
        st.title("AI Research Agent")
        openrouter_client, firecrawl_client = init_clients()

        st.header("Unified Research Interface")

        # --- Add Model Selection Section ---
        st.subheader("Model Selection")
        # Determine the index of the currently configured primary model for default selection
        try:
            default_model_identifier = OPENROUTER_PRIMARY_MODEL # Loaded from config/env
            default_index = MODEL_IDENTIFIERS.index(default_model_identifier)
        except ValueError:
            default_index = 0 # Fallback to the first model if the default isn't in our list
            st.warning(f"Default primary model '{OPENROUTER_PRIMARY_MODEL}' not found in selectable options. Defaulting to first option.")

        selected_model_display_name = st.selectbox(
            "Choose the AI model for report generation:",
            options=MODEL_DISPLAY_NAMES,
            index=default_index,
            key="model_selector",
            help="Select the AI model to use. Defaults are set in config/env."
        )
        # Store the corresponding identifier in session state
        st.session_state.selected_model = MODEL_OPTIONS[selected_model_display_name]
        st.markdown("---") # Separator
        # --- End Model Selection Section ---

        # Research Query Input
        st.subheader("1. Define Your Research Focus")
        research_query = st.text_area(
            "Enter your research query or specific questions:",
            height=100,
            key="research_query_input",
            help="Clearly state what you want the AI to investigate or analyze based on the provided documents and URLs."
        )

        # Document Upload Section
        st.subheader("2. Upload Relevant Documents (Optional)")
        
        uploaded_files_new = st.file_uploader(
            "Upload documents (PDF, DOCX, TXT, MD)",
            type=["pdf", "docx", "txt", "md"],
            accept_multiple_files=True,
            key="unified_file_uploader", # Using a consistent key might be okay if Streamlit handles updates well
            help="Upload whitepapers, reports, articles, or any other relevant documents."
        )

        # Check if the set of uploaded files has actually changed
        current_file_details = [(f.name, f.size) for f in uploaded_files_new] if uploaded_files_new else []
        files_have_changed = (current_file_details != st.session_state.get("last_uploaded_file_details", []))

        if uploaded_files_new and files_have_changed:
            st.session_state.last_uploaded_file_details = current_file_details
            st.session_state.processed_documents_content = [] # Clear previous batch to process anew
            current_batch_processed_content = []
            
            # This inner check is a bit redundant given the outer `if uploaded_files_new` but safe.
            # if uploaded_files_new: 
            with st.status(f"Processing {len(uploaded_files_new)} uploaded file(s)...", expanded=True) as status_container:
                # progress_bar = st.progress(0) # Using st.status handles ongoing updates better for this pattern
                
                for i, uploaded_file_data in enumerate(uploaded_files_new):
                    st.write(f"Processing: {uploaded_file_data.name} ({i+1}/{len(uploaded_files_new)})")
                    # action_details = f"File upload attempt: {uploaded_file_data.name}" # Audit log moved
                    # log_status = "SUCCESS"

                    file_bytes = uploaded_file_data.getvalue()
                    content = ""
                    # Determine file type for processing (prefer extension for clarity)
                    file_extension = uploaded_file_data.name.split('.')[-1].lower()

                    if file_extension == "pdf":
                        content = extract_text_from_pdf(file_bytes)
                    elif file_extension == "docx":
                        content = extract_text_from_docx(file_bytes)
                    elif file_extension in ["txt", "md"]:
                        content = extract_text_from_txt_md(file_bytes)
                    else:
                        st.warning(f"Unsupported file type skipped (extension not matched): {uploaded_file_data.name}")
                        content = "" 
                        # log_status = "SKIPPED_UNSUPPORTED"

                    if content:
                        current_batch_processed_content.append({"name": uploaded_file_data.name, "text": content})
                        st.success(f"Successfully extracted text from: {uploaded_file_data.name}")
                        action_details = f"Successfully extracted content from uploaded file: {uploaded_file_data.name}"
                        log_audit_event(st.session_state.username, st.session_state.get('role','N/A'), "FILE_PROCESS_SUCCESS", action_details)
                    elif file_extension in ["pdf", "docx", "txt", "md"]:
                        st.error(f"Failed to extract text from: {uploaded_file_data.name} (see specific error above if any).")
                        action_details = f"Failed to extract content from uploaded file: {uploaded_file_data.name}"
                        log_audit_event(st.session_state.username, st.session_state.get('role','N/A'), "FILE_PROCESS_FAILURE", action_details)
                    else: # If skipped due to unsupported type and no content
                        action_details = f"Skipped unsupported file type: {uploaded_file_data.name}"
                        log_audit_event(st.session_state.username, st.session_state.get('role','N/A'), "FILE_PROCESS_SKIPPED", action_details)

                    # progress_bar.progress((i + 1) / len(uploaded_files_new)) # Handled by st.status context
                
                st.session_state.processed_documents_content = current_batch_processed_content
                status_container.update(label=f"{len(current_batch_processed_content)} out of {len(uploaded_files_new)} files processed successfully.", state="complete", expanded=False)
                # Rerun to update displays based on newly processed files in session_state
                # This helps ensure UI consistency immediately after processing
                st.rerun() 
        elif not uploaded_files_new: # If no files are selected (uploader is cleared)
            if st.session_state.last_uploaded_file_details: # If there were files before, clear details
                st.session_state.last_uploaded_file_details = []
                st.session_state.processed_documents_content = []
                st.rerun() # Rerun to reflect that no files are processed

        # Display summary of currently processed documents from session_state
        if st.session_state.get("processed_documents_content"):
            st.markdown("---")
            st.subheader(f"Processed Documents ({len(st.session_state.processed_documents_content)} ready for report):")
            for doc_info in st.session_state.processed_documents_content:
                with st.expander(f"{doc_info['name']} ({len(doc_info['text'])} chars) - Click to preview"):
                    st.caption(doc_info['text'][:250] + "..." if len(doc_info['text']) > 250 else doc_info['text'])
            st.markdown("---")

        # URL Input Section
        st.subheader("3. Provide Specific Web URLs (Optional)")
        # st.write("Enter up to 10 URLs for web content scraping:") # Old instruction
        
        # Create a list to hold URL inputs (Old)
        # url_inputs = [] 
        # for i in range(10):
        #     url = st.text_input(f"URL {i+1}", key=f"url_input_{i+1}", placeholder=f"https://example.com/page{i+1}")
        #     url_inputs.append(url)

        # New Text Area for URLs
        urls_text_area = st.text_area(
            "Enter URLs, one per line:",
            height=150, # Adjust height as needed
            key="urls_text_area_input",
            placeholder="https://example.com/page1\nhttps://another-example.com/article2\n..."
        )
        
        # Collect provided URLs from the text area
        # submitted_urls = [url for url in url_inputs if url] # Old collection logic
        if urls_text_area:
            submitted_urls = [url.strip() for url in urls_text_area.split('\n') if url.strip()] 
        else:
            submitted_urls = []

        # --- Add Crawl Input Section --- (Header 4)
        st.subheader("4. Crawl & Scrape Site (Optional)")
        crawl_start_url = st.text_input(
            "Starting URL for Crawl:",
            key="crawl_start_url_input",
            placeholder="https://example.com/startpage"
        )
        crawl_limit = st.number_input(
            "Max Pages to Scrape (Crawl Limit):",
            min_value=1,
            max_value=50, # Set a reasonable upper limit
            value=5, # Default limit
            step=1,
            key="crawl_limit_input",
            help="Maximum number of pages to scrape during the crawl, starting from the URL above."
        )
        st.markdown("---")
        # --- End Crawl Input Section ---

        # Report Generation
        st.subheader("5. Generate Report")
        if st.button("Generate Unified Report", key="generate_unified_report_button"):
            # Modified input check: Need query OR docs OR explicit URLs OR crawl URL
            if not research_query and not st.session_state.processed_documents_content and not submitted_urls and not crawl_start_url:
                st.warning("Please provide a research query, upload documents, enter specific URLs, or provide a starting URL to crawl.")
            else:
                # Clear previous report content and results
                st.session_state.unified_report_content = ""
                st.session_state.scraped_web_content = [] # Clear previous specific scrape results
                st.session_state.crawled_web_content = [] # Clear previous crawl results

                # Log the report generation attempt (include crawl info if provided)
                crawl_info = f" Crawl Start: '{crawl_start_url}', Limit: {crawl_limit}." if crawl_start_url else ""
                processed_doc_names = [doc['name'] for doc in st.session_state.processed_documents_content]
                details_str = f"Research Query: '{research_query}'. Files: {len(processed_doc_names)} ({', '.join(processed_doc_names[:3])}{'...' if len(processed_doc_names) > 3 else ''}). Specific URLs: {len(submitted_urls)}.{crawl_info}"
                
                links_for_log = submitted_urls[:] # Copy list
                if crawl_start_url: links_for_log.append(f"[CRAWL_START] {crawl_start_url}")

                log_audit_event(
                    username=st.session_state.username,
                    role=st.session_state.get('role','N/A'), 
                    action="REPORT_GENERATION_INITIATED",
                    details=details_str,
                    links=links_for_log if links_for_log else None,
                    model_name=st.session_state.get("selected_model", "N/A")
                )

                with st.spinner("Processing inputs and generating report..."):
                    # --- Stage 1a: Scrape Specific URLs ---
                    if submitted_urls:
                        st.info(f"Starting web scraping for {len(submitted_urls)} specific URL(s)...")
                        scraped_data_specific = await process_urls(submitted_urls, firecrawl_client)
                        st.session_state.scraped_web_content = scraped_data_specific # Store specific results
                        # ... (keep existing feedback logic for specific URLs) ...
                    else:
                        st.info("No specific URLs provided for web scraping.")
                        
                    # --- Stage 1b: Crawl & Scrape Site ---
                    if crawl_start_url:
                        # Call the new crawl function
                        crawled_data = await crawl_and_scrape_site(crawl_start_url, crawl_limit, firecrawl_client)
                        st.session_state.crawled_web_content = crawled_data
                        # Feedback is handled within the crawl function
                    else:
                        st.info("No starting URL provided for site crawl.")
                    
                    # --- Task 4: Combined Analysis & Report Generation --- 
                    st.info("Combining processed content and generating AI report...")
                    
                    # 1. Retrieve research query
                    # research_query variable is already available here

                    # 2. Retrieve text from processed documents
                    document_texts = []
                    if st.session_state.processed_documents_content:
                        for doc in st.session_state.processed_documents_content:
                            document_texts.append(f"--- Document: {doc['name']} ---\n{doc['text']}\n---")
                    combined_document_text = "\n".join(document_texts)

                    # 3. Retrieve text from BOTH scraped specific URLs and crawled URLs
                    scraped_content_texts = []
                    if st.session_state.scraped_web_content:
                        scraped_content_texts.append("--- Specific URLs Content ---")
                        for item in st.session_state.scraped_web_content:
                            if item["status"] == "success" and item.get("content"): 
                                scraped_content_texts.append(f"--- URL: {item['url']} ---\n{item['content']}\n---")
                    
                    crawled_content_texts = []
                    successful_crawls = [item for item in st.session_state.crawled_web_content if item["status"] == "success" and item.get("content")]
                    if successful_crawls:
                        crawled_content_texts.append("--- Crawled Site Content ---")
                        for item in successful_crawls:
                            crawled_content_texts.append(f"--- Crawled URL: {item['url']} ---\n{item['content']}\n---")
                                
                    combined_scraped_text = "\n".join(scraped_content_texts)
                    combined_crawled_text = "\n".join(crawled_content_texts)

                    # 4. Combine all text sources with the research query
                    # The research_query acts as the primary instruction/question.
                    # The system_prompt (in st.session_state.system_prompt) guides the AI's persona and output format.
                    
                    # Construct the main instruction based on whether a user query was provided
                    if research_query:
                        full_prompt_for_ai = f"Research Query: {research_query}\\n\\n"
                    else:
                        # Default instruction if query is empty - rely on system prompt and content
                        full_prompt_for_ai = "Research Goal: Please generate a comprehensive report based on the provided content (if any) and the overall objectives defined in the system prompt.\\n\\n"

                    # Append document content if available
                    if combined_document_text:
                        full_prompt_for_ai += f"Provided Document(s) Content:\\n{combined_document_text}\\n\\n"
                    else:
                        full_prompt_for_ai += "No documents were provided or processed.\\n\\n"
                        
                    # Append scraped web content if available
                    if combined_scraped_text:
                        full_prompt_for_ai += f"Provided Specific Web Content:\n{combined_scraped_text}\\n\\n"
                    # else: (No longer needed to say 'No web content' here, check combined below)
                    #    full_prompt_for_ai += "No specific web content was provided or successfully scraped.\\n\\n"
                        
                    if combined_crawled_text:
                        full_prompt_for_ai += f"Provided Crawled Web Content:\n{combined_crawled_text}\\n\\n"
                        
                    if not combined_scraped_text and not combined_crawled_text:
                         full_prompt_for_ai += "No web content was provided or successfully scraped/crawled.\\n\\n"
                    
                    full_prompt_for_ai += "Based on the research goal/query and all the provided content above, please generate a comprehensive report."

                    # 5. Call OpenRouterClient
                    # Ensure openrouter_client is initialized (it is, outside this button block)
                    
                    # --- DEBUGGING --- 
                    st.write("--- DEBUG INFO BEFORE AI CALL ---")
                    # Prepare query snippet for debugging, handling potential quotes
                    debug_query_snippet = research_query[:50]
                    if len(research_query) > 50:
                        debug_query_snippet += "..."
                    st.write(f"Research Query Provided: {'Yes' if research_query else 'No'} (Content: '{debug_query_snippet}')") # Use prepared snippet
                    st.write(f"Combined Document Text Provided: {'Yes' if combined_document_text else 'No'} (Length: {len(combined_document_text)})")
                    st.write(f"Combined Scraped Text Provided: {'Yes' if combined_scraped_text else 'No'} (Length: {len(combined_scraped_text)})")
                    st.write(f"Combined Crawled Text Provided: {'Yes' if combined_crawled_text else 'No'} (Length: {len(combined_crawled_text)})")
                    st.write("--- END DEBUG INFO ---")
                    # --- END DEBUGGING ---

                    try:
                        # --- Get the selected model ---
                        model_to_use = st.session_state.get("selected_model", OPENROUTER_PRIMARY_MODEL) # Fallback just in case
                        # Remove the log here as it's logged earlier and below
                        # log_audit_event(st.session_state.username, st.session_state.get('role','N/A'), "MODEL_SELECTION_USED", f"Using model: {model_to_use} for report generation.")
                        # --- End Get model ---

                        ai_generated_report = await openrouter_client.generate_response(
                            prompt=full_prompt_for_ai, 
                            system_prompt=st.session_state.system_prompt,
                            model_override=model_to_use # Pass selected model
                        )
                        if ai_generated_report:
                            st.session_state.unified_report_content = ai_generated_report
                            st.success("AI Report generated successfully!")
                            
                            successfully_scraped_urls = [item['url'] for item in st.session_state.scraped_web_content if item["status"] == "success"]
                            query_for_log = research_query if research_query else '[SYSTEM PROMPT]'
                            success_details = f"AI report generated for query: '{query_for_log}'. Docs: {len(processed_doc_names)}. Scraped URLs: {len(successfully_scraped_urls)}."
                            
                            log_audit_event(
                                username=st.session_state.username,
                                role=st.session_state.get('role','N/A'), 
                                action="REPORT_GENERATION_SUCCESS", 
                                details=success_details,
                                links=successfully_scraped_urls if successfully_scraped_urls else None,
                                model_name=model_to_use # Log model used
                            )
                        else:
                            st.session_state.unified_report_content = "Failed to generate AI report. The AI returned an empty response."
                            st.error("AI report generation failed or returned empty.")
                            query_for_log = research_query if research_query else '[SYSTEM PROMPT]'
                            log_audit_event(username=st.session_state.username, role=st.session_state.get('role','N/A'), action="REPORT_GENERATION_FAILURE", details=f"AI returned empty report for query: '{query_for_log}'")
                            log_audit_event(username=st.session_state.username, role=st.session_state.get('role','N/A'), action="REPORT_GENERATION_FAILURE", details=f"AI returned empty report for query: '{query_for_log}'", model_name=model_to_use)
                    except Exception as e:
                        st.session_state.unified_report_content = f"An error occurred during AI report generation: {str(e)}"
                        st.error(f"Error calling AI: {e}")
                        query_for_log = research_query if research_query else '[SYSTEM PROMPT]'
                        log_audit_event(username=st.session_state.username, role=st.session_state.get('role','N/A'), action="REPORT_GENERATION_ERROR", details=f"Error during AI call for query '{query_for_log}': {str(e)}")
                        log_audit_event(username=st.session_state.username, role=st.session_state.get('role','N/A'), action="REPORT_GENERATION_ERROR", details=f"Error during AI call for query '{query_for_log}': {str(e)}", model_name=model_to_use)
                    
                    # Remove the debug JSON output now that real processing is in place
                    # st.markdown("--- DEBUG: Information Collected ---\")
                    # st.json({
                    #     "research_query": research_query,
                    #     "processed_documents": st.session_state.processed_documents_content,
                    #     "scraped_urls_content": st.session_state.scraped_web_content
                    # })
                    # st.info("Report generation logic (Task 4) is pending. Displaying collected data for now.")
                    # --- End of Task 4 ---
                
                # Rerun to display the generated report or messages
                st.rerun()
        
        # Area for actual report display (will be populated by backend logic)
        if "unified_report_content" in st.session_state and st.session_state.unified_report_content:
            st.markdown("---")
            st.subheader("Generated Report")
            st.markdown(st.session_state.unified_report_content)
            st.download_button(
                label="Download Unified Report",
                data=st.session_state.unified_report_content,
                file_name="unified_research_report.md",
                mime="text/markdown",
                key="download_actual_unified_report",
                on_click=lambda: log_audit_event(
                    username=st.session_state.username, 
                    role=st.session_state.get('role', 'N/A'), 
                    action="REPORT_DOWNLOADED", 
                    details=f"User downloaded report: unified_research_report.md for query '{(st.session_state.get('research_query_input', 'QUERY_NOT_IN_SESSION_FOR_DOWNLOAD_LOG'))}'"
                )
            )
        # elif "generate_unified_report_button_clicked" in st.session_state and \
        #      st.session_state.generate_unified_report_button_clicked and \
        #      not st.session_state.unified_report_content:
        #      # This logic might be implicitly handled by the spinner and rerun, 
        #      # or if a specific message for "no report generated" is desired, it can be placed here.
        #      # For now, if content is empty, nothing is shown, which is acceptable.
        #      pass 

        st.markdown("---") # Separator before Admin Panel
        
        # ==== ADMIN PANEL ====
        if st.session_state.get("role") == "admin":
            st.header("Admin Panel - Audit Logs")
            
            log_file_path = Path("/app/logs/audit.log")
            log_data = []

            if log_file_path.exists():
                try:
                    with open(log_file_path, 'r') as f:
                        # Read lines in reverse to show newest first
                        log_lines = reversed(f.readlines())
                        for line in log_lines:
                            parsed = parse_log_line(line)
                            if parsed:
                                log_data.append(parsed)
                                
                    if log_data:
                        df = pd.DataFrame(log_data)
                        # Reorder columns for better readability
                        cols_order = ["Timestamp", "Username", "Role", "Action", "Model", "Details", "Links"]
                        cols_to_display = [col for col in cols_order if col in df.columns]
                        
                        # Configure column widths
                        column_config = {
                            "Links": st.column_config.TextColumn("Links", width="medium"),
                            "Details": st.column_config.TextColumn("Details", width="medium"),
                             # Set timestamp width if needed
                            "Timestamp": st.column_config.TextColumn("Timestamp", width="small"),
                        }
                        # Apply config only for columns present in the dataframe
                        active_column_config = {k: v for k, v in column_config.items() if k in cols_to_display}
                        
                        st.dataframe(df[cols_to_display], column_config=active_column_config)
                    else:
                        st.info("Audit log file exists but contains no parseable entries.")
                        
                except Exception as e:
                    st.error(f"Error reading or parsing audit log file: {e}")
            else:
                st.warning("Audit log file not found. Logging may not be configured or no events logged yet.")
            
            if st.button("Refresh Logs", key="refresh_logs_button"):
                 st.rerun() # Simple way to refresh the view by rerunning the script
        # ==== END ADMIN PANEL ====

    else: # Not authenticated
        st.info("Please log in or create an account to access the research pipeline.")
        # The login/signup logic from the sidebar already handles this visibility.
        # This message is a fallback or main page content for unauthenticated users.

async def run_main():
    await main()

if __name__ == "__main__":
    asyncio.run(run_main()) 